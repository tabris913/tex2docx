# coding: utf-8

import collections
import os
from pathlib import Path
import uuid
from typing import List, Union

import docx
import matplotlib.mathtext as mathtext
import pdf2image

from ..tex_parser.document import Document as TexDoc
from ..tex_parser.document import *
from ..tex_parser.document.command.section import *
from ..tex_parser.document.element import TexElement
from ..tex_parser.document.utils import Children
from .ref import Reference

WIDTH = docx.shared.Cm(15)


class Document(docx.document.Document):
    def __init__(self, source: Path):
        self.__source_dir = source.parent
        self.__base = docx.Document()
        super().__init__(self.__base.element, self.__base.part)
        self.__used_imgs = []
        self.__refs: List[Reference] = []
        self.__tags = collections.defaultdict(lambda: 'unknown ref')
        self.__table_count = 0
        self.__figure_count = 0
        self.__code_count = 0

    @property
    def used_imgs(self) -> List[str]:
        return self.__used_imgs

    def add_title(self, title: str):
        self.add_heading(title, 0)

    def add_author(self, author: str):
        author = self.add_paragraph(author)
        author.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT

    def add_date(self, date: str):
        date = self.add_paragraph(date)
        date.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT

    def parse_document(self, document: Union[TexDoc, TexElement]):
        for child in document.children:
            if isinstance(child, str):
                ...  # print(child)
            elif child.children is not None:
                # env, section
                if isinstance(child, SectionBase):
                    if isinstance(child, Section):
                        self.add_section(
                            docx.enum.section.WD_SECTION_START.CONTINUOUS)
                        self.add_heading(child.name, 1)
                    elif isinstance(child, SubSection):
                        self.add_heading(child.name, 2)
                    elif isinstance(child, SubSubSection):
                        self.add_heading(child.name, 3)
                    elif isinstance(child, Paragraph):
                        self.add_heading(child.name, 4)
                    elif isinstance(child, SubParagraph):
                        self.add_heading(child.name, 5)
                    else:
                        ...  # print(child)
                elif isinstance(child, Table):
                    self.__add_table(child)
                    continue
                elif isinstance(child, Figure):
                    self.__add_figure(child)
                    continue
                self.parse_document(child)
            elif isinstance(child, ClearPage):
                self.add_page_break()
            elif isinstance(child, ClearDoublePage):
                self.add_page_break()
                self.add_page_break()
            elif isinstance(child, Text):
                self.add_paragraph(child.text)
            elif isinstance(child, TableOfContents):
                self.add_paragraph('[toc]')
            elif isinstance(child, Formula):
                self.__add_formula(child)
            elif isinstance(child, Ref):
                self.__refs.append(
                    Reference(child.parameter, self.add_paragraph()))
            elif isinstance(child, Label):
                if isinstance(document, SectionBase):
                    self.__tags[child.parameter] = document.label
                else:
                    self.__tags[child.paragraph] = f'Label: {child.paragraph}'
            else:
                self.add_paragraph(str(child))
                # print(child)

    def __add_table(self, table: Table):
        if table.has_tabular:
            self.__table_count += 1

            p = self.add_paragraph()
            p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

            # add caption if exists
            try:
                caption = table.caption
                p.add_run(
                    f'Table {self.__table_count}: {caption.parameter}').bold = True
                label = table.label
                self.__tags[label.parameter] = f'Table {self.__table_count}'
            except BaseException:
                print('failed')

            # add table
            tab = self.add_table(*table.shape)
            for child in table.children:
                if isinstance(child, Tabular):
                    for i in range(child.rows):
                        for j in range(child.cols):
                            tab.rows[i].cells[j].text = child.table[i][j]
                elif isinstance(child, Centering):
                    tab.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            self.add_paragraph()

    def __add_figure(self, figure: Figure):
        if figure.has_graphics:
            self.__figure_count += 1

            # add figure
            for child in figure.children:
                if isinstance(child, IncludeGraphics):
                    path = self.__source_dir / Path(child.parameter)
                    if path.exists():
                        if path.suffix == '.pdf':
                            try:
                                img = pdf2image.convert_from_path(str(path))[0]
                                newpath = (
                                    path.parent / path.stem).with_suffix('.png')
                                img.save(newpath, 'png')
                                fig = self.add_picture(str(newpath))
                            except BaseException as e:
                                raise Exception(
                                    f'occured while converting pdf to png: {e}')
                        else:
                            fig = self.add_picture(str(path))
                        if fig.width > WIDTH:
                            ratio = fig.width / WIDTH
                            fig.width = WIDTH
                            fig.height = int(fig.height / ratio)
                    else:
                        raise FileNotFoundError(path)

            p = self.add_paragraph()
            p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

            # add caption if exists
            try:
                caption = figure.caption
                p.add_run(
                    f'Figure {self.__figure_count}: {caption.parameter}').bold = True
                label = figure.label
                self.__tags[label.parameter] = f'Figure {self.__figure_count}'
            except BaseException:
                print('failed')

            self.add_paragraph()

    def __add_formula(self, formula: Formula):
        figname = f'fig/{uuid.uuid4()}.png'
        p = self.add_paragraph()
        p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run()
        parser = mathtext.MathTextParser('bitmap')
        offset = parser.to_png(figname, formula.text)
        fig = r.add_picture(figname)
        self.__used_imgs.append(figname)

    def assign_refs(self):
        for ref in self.__refs:
            ref.paragraph.add_run(self.__tags[ref.tag])
